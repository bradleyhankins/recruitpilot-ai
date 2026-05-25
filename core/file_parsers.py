from __future__ import annotations

from io import BytesIO


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""

    reader = PdfReader(BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text.strip())
    return "\n\n".join(pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""

    document = Document(BytesIO(file_bytes))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    table_text = []
    for table in document.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_values:
                table_text.append(" | ".join(row_values))

    return "\n".join(paragraphs + table_text)


def extract_text_from_plain_text(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore")


def extract_resume_text(file_bytes: bytes, file_name: str) -> tuple[str, str | None]:
    extension = file_name.lower().rsplit(".", 1)[-1] if "." in file_name else ""

    if extension == "pdf":
        text = extract_text_from_pdf(file_bytes)
        if not text.strip():
            return "", "No readable text was found in this PDF. It may be scanned or image-based. Please paste the resume text manually or upload a text-based PDF/DOCX."
        return text, None

    if extension == "docx":
        text = extract_text_from_docx(file_bytes)
        if not text.strip():
            return "", "No readable text was found in this Word document. Please paste the resume text manually or upload another file."
        return text, None

    if extension in {"txt", "md"}:
        return extract_text_from_plain_text(file_bytes), None

    if extension == "doc":
        return "", "Legacy .doc files are not supported in this demo. Please save the document as .docx or PDF and upload again."

    return "", "Unsupported file type. Please upload a PDF, DOCX, TXT, or MD file."
